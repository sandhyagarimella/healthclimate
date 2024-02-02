# Text Processing

import os
import re
from collections import Counter
from docx import Document
import spacy
import pandas as pd
import nltk
from nltk.tokenize import sent_tokenize
from wordcloud import WordCloud
import matplotlib.pyplot as plt
from heapq import nlargest


path = r'/Users/sandhyagarimella/Documents/GitHub/final-project-climate-and-health/Data'
doc_name = 'CHINA_Climate_Strategy_2035.docx'
doc_path = os.path.join(path, doc_name)

doc_name_eu = 'EU_202020_Climate_Policy_2008.docx'
doc_path_eu = os.path.join(path, doc_name_eu)

###### Parsing the word document - China######
def process_docx(doc):
    """ Process the DOCX file, merging paragraphs, including tables, and removing specific texts.
    Excludes texts that are likely titles, footnotes, or citations (less than 312 characters),
    as well as texts with exactly 769 characters. """
    processed_paragraphs = []
    merge_text = ""

    # Process paragraphs
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()

        # Exclude specific texts based on length
        if len(text) < 312 or len(text) == 769:
            continue

        # Merge with previous paragraph if empty (indicating page break)
        if text:
            if merge_text:
                processed_paragraphs.append(merge_text + " " + text)
                merge_text = ""
            else:
                processed_paragraphs.append(text)
        else:
            merge_text = processed_paragraphs.pop() if processed_paragraphs else ""

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text and (len(cell_text) >= 312 and len(cell_text) != 769):
                    processed_paragraphs.append(cell_text)

    return processed_paragraphs

# Keyword and Sentence extraction function
def find_health_related_sentences(paragraphs, keywords):
    """Find sentences with specific health-related keywords and include one sentence before and after."""
    relevant_sentences = []
    
    for paragraph in paragraphs:
        # Split paragraph into sentences
        sentences = sent_tokenize(paragraph)

        for i, sentence in enumerate(sentences):
            # Check if the sentence contains any of the health-related keywords
            if any(keyword in sentence.lower() for keyword in keywords):
                # Add previous sentence if it exists --- Cite ChatGPT
                if i > 0:
                    relevant_sentences.append(sentences[i - 1])

                # Add the current sentence
                relevant_sentences.append(sentence)

                # Add next sentence if it exists
                if i < len(sentences) - 1:
                    relevant_sentences.append(sentences[i + 1])

    return relevant_sentences


# Function to summarize paragraphs using spaCy --- cite ChatGPT for example code adapted here manually
def spacy_summarize_paragraph(paragraph):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(paragraph)
    keywords = [token.text for token in doc if not token.is_stop and not token.is_punct]
    freq_word = Counter(keywords)
    max_freq = freq_word.most_common(1)[0][1]
    for word in freq_word.keys():
        freq_word[word] = (freq_word[word] / max_freq)
    sent_strength = {}
    for sent in doc.sents:
        for word in sent:
            if word.text in freq_word.keys():
                if sent in sent_strength.keys():
                    sent_strength[sent] += freq_word[word.text]
                else:
                    sent_strength[sent] = freq_word[word.text]
    summary_sentences = nlargest(3, sent_strength, key=sent_strength.get)
    return ' '.join([str(sent) for sent in summary_sentences])

# Function to extract entities for word cloud
def extract_entities(text):
    nlp = spacy.load("en_core_web_sm")
    doc = nlp(text)
    return ' '.join([ent.text for ent in doc.ents])

########################## EU Policy Document ############################

#Parsing EU Policy 
def load_document_paragraphs(file_path):

    try:
        # Load the document
        doc = Document(file_path)

        # Filter paragraphs based on conditions - removing titles generally less than 62 characters and in uppercase
        paragraphs = [p.text for p in doc.paragraphs 
                      if p.text.strip() != '' and len(p.text.strip()) >= 62 and not p.text.isupper()]

        return paragraphs
    except Exception as e:
        return f"An error occurred: {e}"

############################################################################

# Main execution- China
doc = Document(doc_path)
final_paragraphs = process_docx(doc)
health_keywords = [
    'health risk', 'pollution', 'disease', 'diseases', 'illness', 
    'infection', 'infectious', 'infections', 'asthma', 'public health', 
    'wellbeing', 'hospital', 'virus', 'chronic', 'vaccine', 'treatment', 
    'medication', 'medicine'
]
health_related_sentences = find_health_related_sentences(final_paragraphs, health_keywords)

# Summarize each paragraph with health-related sentences
summaries = [spacy_summarize_paragraph(paragraph) for paragraph in health_related_sentences]

# Concatenate summaries into a single string, separating them with a unique delimiter
concatenated_summaries = ' || '.join(summaries)

# Generate tag cloud from summmaries
tag_cloud = WordCloud(width=800, height=800, background_color='white').generate(concatenated_summaries)

# Display tag cloud
plt.figure(figsize=(8, 8), facecolor=None)
plt.imshow(tag_cloud)
plt.axis("off")
plt.tight_layout(pad=0)
plt.show()

# Extract entities and create word cloud
all_text = ' '.join(health_related_sentences)
entities_text = extract_entities(all_text)
wordcloud = WordCloud(width=800, height=800, background_color='white').generate(entities_text)

# Display word cloud
plt.figure(figsize=(8, 8), facecolor=None)
plt.imshow(wordcloud)
plt.axis("off")
plt.tight_layout(pad=0)
plt.show()

# Execution - EU
eu_paragraphs = load_document_paragraphs(doc_path_eu)

# Extract health-related sentences from the EU document
eu_health_related_sentences = find_health_related_sentences(eu_paragraphs, health_keywords)

# Summarize each paragraph with health-related sentences in the EU document
eu_summaries = [spacy_summarize_paragraph(paragraph) for paragraph in eu_health_related_sentences]

# Concatenate EU summaries into a single string
eu_concatenated_summaries = ' || '.join(eu_summaries)

# Generate and display tag cloud for EU summaries
eu_tag_cloud = WordCloud(width=800, height=800, background_color='white').generate(eu_concatenated_summaries)
plt.figure(figsize=(8, 8), facecolor=None)
plt.imshow(eu_tag_cloud)
plt.axis("off")
plt.tight_layout(pad=0)
plt.show()

# Extract entities from the EU health-related sentences and create word cloud
eu_all_text = ' '.join(eu_health_related_sentences)
eu_entities_text = extract_entities(eu_all_text)
eu_wordcloud = WordCloud(width=800, height=800, background_color='white').generate(eu_entities_text)

# No entities found for EU